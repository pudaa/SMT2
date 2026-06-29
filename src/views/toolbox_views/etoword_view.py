import sys
from PySide6.QtWidgets import (QWidget, QVBoxLayout, QHBoxLayout, QLabel, 
                               QSplitter, QComboBox, QListWidget, QLineEdit,
                               QScrollArea, QPushButton, QApplication,
                               QMessageBox)
from PySide6.QtCore import Qt, QMimeData, QByteArray, Signal, QObject, QThread, QTimer
from PySide6.QtGui import QDragEnterEvent, QDropEvent, QDragMoveEvent, QDrag
from PySide6.QtWidgets import QProgressDialog
from typing import Optional
from collections import OrderedDict
from copy import deepcopy

class DraggableListWidget(QListWidget):
    """可拖拽的列表组件"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setDragEnabled(True)
        self.setAcceptDrops(True)
        self.setDragDropMode(QListWidget.DragDropMode.DragOnly)
    
    def startDrag(self, supportedActions):
        """开始拖拽操作"""
        item = self.currentItem()
        if item:
            mimeData = QMimeData()
            # 获取该列表自身的项目类型（由 update_title_list 设置），优先使用列表自身的 item_mode
            mode = getattr(self, 'item_mode', None)
            if mode is None:
                # 回退到父视图的模式下拉
                mode = "列模式"
                parent = self.parent()
                if parent and hasattr(parent, 'mode_combo') and parent.mode_combo is not None:
                    try:
                        mode = parent.mode_combo.currentText()
                    except Exception:
                        mode = "列模式"
            # 将数据序列化为字节流
            data_str = f"{item.text()}|{mode}"
            mimeData.setData("application/x-etoword-item", QByteArray(data_str.encode()))
            
            drag = QDrag(self)
            drag.setMimeData(mimeData)
            drag.exec_(Qt.DropAction.CopyAction)

class DropAreaWidget(QWidget):
    """可接收拖拽的区域"""
    item_dropped = Signal(str, str)  # 发送标题和模式信号
    
    def __init__(self):
        super().__init__()
        self.setAcceptDrops(True)
        self.setup_ui()
    
    def setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(10, 10, 10, 10)
        self.label = QLabel("拖拽项目到这里添加到模板")
        self.label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.label.setObjectName("dropAreaLabel")
        layout.addWidget(self.label)
    
    def dragEnterEvent(self, event):
        if event.mimeData().hasFormat("application/x-etoword-item"):
            event.acceptProposedAction()
        else:
            event.ignore()
    
    def dropEvent(self, event):
        if event.mimeData().hasFormat("application/x-etoword-item"):
            data = event.mimeData().data("application/x-etoword-item").data().decode()
            title, mode = data.split("|")
            self.item_dropped.emit(title, mode)
            event.acceptProposedAction()
        else:
            event.ignore()

        

class PreviewItemWidget(QWidget):
    """预览项小部件 - 同时支持显示和拖拽"""
    # 定义删除信号
    item_deleted = Signal(int)  # 发送要删除的索引
    
    def __init__(self, index, item, parent=None):
        super().__init__(parent)
        self.index = index
        self.item = item
        self.setAcceptDrops(True)
        self.setup_ui()
    
    def setup_ui(self):
        layout = QHBoxLayout(self)
        layout.setContentsMargins(2, 2, 2, 2)
        layout.setSpacing(5)
        
        # 显示标签（可点击编辑）
        self.label = QLabel(self.item["value"])
        self.label.setWordWrap(True)
        self.label.setObjectName("previewItemLabel")
        layout.addWidget(self.label)
        
        # 样式选择
        self.style_combo = QComboBox()
        self.style_combo.addItems([
            "正文",
            "一级标题",
            "二级标题",
            "三级标题",
            "加粗",
            "斜体",
            "引用",
            "无序列表",
            "有序列表",
        ])
        self.style_combo.setCurrentText(self.item["style"])
        self.style_combo.currentTextChanged.connect(self.on_style_changed)
        self.style_combo.setFixedWidth(100)
        layout.addWidget(self.style_combo)
        
        # 删除按钮
        self.delete_btn = QPushButton("✕")
        self.delete_btn.setFixedSize(40, 25)
        self.delete_btn.clicked.connect(self.on_delete_clicked)
        layout.addWidget(self.delete_btn)
        
        # 拖拽支持
        self.setMouseTracking(True)
    
    def on_style_changed(self, style):
        self.item["style"] = style
        # 更新显示文本
        self.label.setText(self.item["value"])
    
    def on_delete_clicked(self):
        """发送删除信号给父视图，由父视图负责从数据模型中删除并刷新预览"""
        self.item_deleted.emit(self.index)
    
    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self.start_drag()
        try:
            super().mousePressEvent(event)
        except RuntimeError:
            # 有时父组件在拖拽期间已删除该C++对象，捕获并忽略该错误以避免崩溃
            pass
    
    def start_drag(self):
        """开始拖拽操作"""
        mimeData = QMimeData()
        data_str = f"{self.item['title']}|preview_item"
        mimeData.setData("application/x-etoword-item", QByteArray(data_str.encode()))
        
        drag = QDrag(self)
        drag.setMimeData(mimeData)
        drag.exec_(Qt.DropAction.MoveAction)
    
    def dragEnterEvent(self, event):
        if event.mimeData().hasFormat("application/x-etoword-item"):
            event.acceptProposedAction()
        else:
            super().dragEnterEvent(event)
    
    def dropEvent(self, event):
        if event.mimeData().hasFormat("application/x-etoword-item"):
            data = event.mimeData().data("application/x-etoword-item").data().decode()
            title, mode = data.split("|")
            
            # 如果是从左侧列表拖来的，添加新项目
            if mode in ["列模式", "行模式"]:
                parent = self.parent()
                if parent and hasattr(parent, 'add_item_to_template_at_position'):
                    parent.add_item_to_template_at_position(title, mode, self.index + 1)
            
            event.acceptProposedAction()
        else:
            super().dropEvent(event)

class CompositePreviewItemWidget(QWidget):
    """组合预览项小部件 - 显示多个字段项，支持拖拽添加字段项"""
    class ColumnPreviewWidget(QWidget):
        """表示一个列容器：显示列标题以及该列中横向排列的字段项，支持拖拽向该列添加字段"""
        def __init__(self, column_title, items, parent=None):
            super().__init__(parent)
            self.column_title = column_title
            self.items = items  # 列内的字段项列表
            self.setMaximumHeight(88)
            self.setMinimumHeight(88)
            self.setAcceptDrops(True)
            self.setup_ui()

        def setup_ui(self):
            layout = QVBoxLayout(self)
            layout.setContentsMargins(4, 4, 4, 4)
            layout.setSpacing(4)

            self.title_label = QLabel(f"{self.column_title}")
            self.title_label.setObjectName("columnTitle")
            layout.addWidget(self.title_label)

            # 容纳列内字段的横向容器
            self.items_container = QWidget()
            self.items_container.setObjectName("columnItemsContainer")
            self.items_layout = QHBoxLayout(self.items_container)
            self.items_layout.setContentsMargins(0, 0, 0, 0)
            self.items_layout.setSpacing(6)
            layout.addWidget(self.items_container)

            # 用 objectName 标识列容器以便 QSS 控制样式
            self.setObjectName("columnWidget")

            self.refresh_items()

        def refresh_items(self):
            # 清除旧项
            while self.items_layout.count():
                it = self.items_layout.takeAt(0)
                if it and it.widget():
                    it.widget().setParent(None)

            # 添加每个字段为小标签（可删除、可拖拽的 PreviewItemWidget）
            for i, item in enumerate(self.items):
                w = PreviewItemWidget(i, item, self.items_container)
                # 连接删除到父控件的接口（父对象是 ExcelToWordView）
                w.item_deleted.connect(self.handle_item_deleted)
                self.items_layout.addWidget(w)

        def handle_item_deleted(self, index):
            # 删除列内项并请求父刷新
            if 0 <= index < len(self.items):
                self.items.pop(index)
                self.refresh_items()
                parent = self.parent()
                # 向上寻找 ExcelToWordView 并让其刷新整体预览
                while parent and not hasattr(parent, 'refresh_preview_area'):
                    parent = parent.parent()
                # 如果当前列已无项，则从父的 preview_columns 中移除该列
                if not self.items:
                    container = parent
                    while container and not hasattr(container, 'preview_columns'):
                        container = container.parent()
                    if container and self.column_title in container.preview_columns:
                        try:
                            del container.preview_columns[self.column_title]
                        except KeyError:
                            pass
                if parent:
                    parent.refresh_preview_area()

        def dragEnterEvent(self, event):
            if event.mimeData().hasFormat("application/x-etoword-item"):
                event.acceptProposedAction()
            else:
                super().dragEnterEvent(event)

        def dropEvent(self, event):
            if event.mimeData().hasFormat("application/x-etoword-item"): # 列内字段拖拽
                data = event.mimeData().data("application/x-etoword-item").data().decode()
                title, mode = data.split("|")
                # 调用父控件的方法把字段加入到该列
                parent = self.parent()
                while parent and not hasattr(parent, 'add_item_to_column'):
                    parent = parent.parent()
                if parent:
                    parent.add_item_to_column(self.column_title, title, mode)
                event.acceptProposedAction()
            else:
                super().dropEvent(event)

    """复合预览项小部件 - 支持在同一行显示多个字段"""
    def __init__(self, index, items, parent=None):
        super().__init__(parent)
        self.index = index
        self.items = items  # 包含多个item的列表
        self.setAcceptDrops(True)
        self.setup_ui()
    
    def setup_ui(self):
        layout = QHBoxLayout(self)
        layout.setContentsMargins(2, 2, 2, 2)
        layout.setSpacing(5)
        
        # 显示组合内容
        self.content_label = QLabel()
        self.update_display_text()
        self.content_label.setWordWrap(True)
        self.content_label.setObjectName("compositeContentLabel")
        # layout.addWidget(self.content_label)
        
        
        # 拖拽支持
        self.setMouseTracking(True)
    
    def update_display_text(self):
        """更新显示文本"""
        if not self.items:
            self.content_label.setText("空组合项")
            return
            
        # 组合显示所有字段的值
        display_parts = []
        for i, item in enumerate(self.items):
            value = item.get("value", "")
            if value:
                display_parts.append(f"{value}")
            else:
                display_parts.append("")
        
        display_text = " ".join(display_parts)
        self.content_label.setText(display_text)
    
    def on_style_changed(self, style):
        # 更新所有子项目样式
        for item in self.items:
            item["style"] = style
        self.update_display_text()
    
    def add_field_to_composite(self):
        """向组合项添加新字段"""
        # 这里可以通过弹窗或其他方式让用户选择要添加的字段
        # 暂时留空，后续实现具体的添加逻辑
        pass
    
    def delete_item(self):
        parent = self.parent()
        if parent and hasattr(parent, 'remove_composite_preview_item'):
            parent.remove_composite_preview_item(self.index)
    
    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self.start_drag()
        try:
            super().mousePressEvent(event)
        except RuntimeError:
            # 忽略在拖拽期间C++对象被删除导致的错误
            pass
    
    def start_drag(self):
        """开始拖拽操作"""
        mimeData = QMimeData()
        # 将组合项标识为特殊类型
        titles = "|".join([item["title"] for item in self.items])
        data_str = f"{titles}|composite_item"
        mimeData.setData("application/x-etoword-item", QByteArray(data_str.encode()))
        
        drag = QDrag(self)
        drag.setMimeData(mimeData)
        drag.exec_(Qt.DropAction.MoveAction)
    
    def dragEnterEvent(self, event):
        if event.mimeData().hasFormat("application/x-etoword-item"):
            event.acceptProposedAction()
        else:
            super().dragEnterEvent(event)
    
    def dropEvent(self, event):
        if event.mimeData().hasFormat("application/x-etoword-item"):
            data = event.mimeData().data("application/x-etoword-item").data().decode()
            title, mode = data.split("|", 1)
            
            # 如果是普通字段，添加到组合中
            if mode in ["列模式", "行模式"]:
                self.add_field_to_existing(title, mode)
            # 如果是另一个组合项，合并
            elif mode == "composite_item":
                self.merge_with_composite(title)
            
            event.acceptProposedAction()
        else:
            super().dropEvent(event)
    
    def add_field_to_existing(self, title, mode):
        """向现有组合添加字段"""
        parent = self.parent()
        if parent and hasattr(parent, 'excel_data'):
            excel_data = parent.excel_data
            if excel_data is not None:
                # 获取字段值
                if mode == "列模式" and title in excel_data.columns:
                    value = str(excel_data.iloc[0][title]) if len(excel_data) > 0 else ""
                elif mode == "行模式":
                    try:
                        if title.startswith("行 "):
                            row_index = int(title.split()[1])
                        else:
                            row_index = int(title)
                        if 0 <= row_index < len(excel_data):
                            value = str(excel_data.iloc[row_index].tolist())
                        else:
                            value = ""
                    except (ValueError, IndexError):
                        value = ""
                else:
                    value = ""
                
                # 添加到组合
                new_item = {"title": title, "value": value, "style": "正文"}
                self.items.append(new_item)
                self.update_display_text()
    
    def merge_with_composite(self, titles_string):
        """与其他组合项合并"""
        # 解析要合并的标题
        titles_to_merge = titles_string.split("|")
        
        parent = self.parent()
        if parent and hasattr(parent, 'excel_data'):
            excel_data = parent.excel_data
            if excel_data is not None:
                # 为每个标题创建项目并添加到当前组合
                for title in titles_to_merge:
                    if title in excel_data.columns:
                        value = str(excel_data.iloc[0][title]) if len(excel_data) > 0 else ""
                        new_item = {"title": title, "value": value, "style": "正文"}
                        self.items.append(new_item)
                
                self.update_display_text()

class ExcelToWordView(QWidget):
    def __init__(self):
        super().__init__()
        self.excel_data: Optional["pd.DataFrame"] = None
        self.preview_columns = OrderedDict()  # 列模式：每列包含多个横向字段
        self.setAcceptDrops(True)  # 启用拖拽功能
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()
        layout.setContentsMargins(0, 0, 0, 0)  # 减少外边距
        layout.setSpacing(5)  # 减少元素间距
        
        # 初始状态：居中QLabel提示
        self.initial_label = QLabel("拖入Excel文件...")
        self.initial_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.initial_label)
        
        # 主要功能区域（初始隐藏）
        self.main_widget = QWidget()
        self.main_layout = QHBoxLayout(self.main_widget)
        self.main_layout.setContentsMargins(0, 0, 0, 0)
        self.main_layout.setSpacing(5)
        
        # 左侧区域
        left_widget = QWidget()
        left_widget.setFixedWidth(120)
        left_layout = QVBoxLayout(left_widget)
        left_layout.setContentsMargins(0, 0, 5, 0)
        left_layout.setSpacing(5)
        
        # 恢复模式下拉（行/列），以便用户选择使用行索引还是列名
        mode_row = QWidget()
        mode_row_layout = QHBoxLayout(mode_row)
        mode_row_layout.setContentsMargins(0, 0, 0, 0)
        mode_row_layout.setSpacing(4)
        mode_row_layout.addWidget(QLabel("模式:"))
        self.mode_combo = QComboBox()
        self.mode_combo.addItems(["列模式", "行模式"])
        self.mode_combo.setCurrentText("列模式")
        self.mode_combo.currentTextChanged.connect(self.on_mode_changed)
        mode_row_layout.addWidget(self.mode_combo)
        left_layout.addWidget(mode_row)

        self.title_list = DraggableListWidget()
        self.title_list.setParent(self)
        left_layout.addWidget(self.title_list)
        
        # 右侧区域
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.setContentsMargins(5, 0, 5, 0)
        right_layout.setSpacing(5)
        
        self.primary_title = QLineEdit()
        self.primary_title.setPlaceholderText("一级标题")
        right_layout.addWidget(self.primary_title)
        
        # 布局方向控件已移除（用户要求删除）
        
        # 整合的预览和拖拽区域
        self.preview_area = QScrollArea()
        self.preview_widget = QWidget()
        self.preview_widget.setObjectName("previewWidget")
        self.preview_layout = QVBoxLayout(self.preview_widget)  # 默认纵向
        self.preview_layout.setContentsMargins(5, 5, 5, 5)
        self.preview_layout.setSpacing(3)
        self.preview_area.setWidget(self.preview_widget)
        self.preview_area.setWidgetResizable(True)
        self.preview_area.setAcceptDrops(True)
        self.preview_area.dragEnterEvent = self.preview_drag_enter
        self.preview_area.dropEvent = self.preview_drop_event
        right_layout.addWidget(self.preview_area)
        
        # 在预览区下方添加空白拖拽区以延伸底部 20px
        self.bottom_drop_zone = DropAreaWidget()
        self.bottom_drop_zone.setFixedHeight(20)
        # 连接到底部放下的处理器
        self.bottom_drop_zone.item_dropped.connect(self.handle_bottom_drop)
        # 设置透明背景以不影响视觉（通过QSS保持最小冲突）
        right_layout.addWidget(self.bottom_drop_zone)
        
        self.generate_btn = QPushButton("生成Word")
        self.generate_btn.clicked.connect(self.generate_word)
        right_layout.addWidget(self.generate_btn)
        
        # 分割器
        splitter = QSplitter(Qt.Horizontal)
        splitter.addWidget(left_widget)
        splitter.addWidget(right_widget)
        splitter.setSizes([200, 400])  # 调整默认比例
        self.main_layout.addWidget(splitter)
        
        layout.addWidget(self.main_widget)
        self.main_widget.hide()
        
        self.setLayout(layout)

    def dragEnterEvent(self, event: QDragEnterEvent):
        if event.mimeData().hasUrls():
            # 检查是否为Excel文件
            urls = event.mimeData().urls()
            if urls:
                file_path = urls[0].toLocalFile()
                if file_path.endswith(('.xlsx', '.xls')):
                    event.acceptProposedAction()
                else:
                    event.ignore()
        else:
            event.ignore()

    def dropEvent(self, event: QDropEvent):
        urls = event.mimeData().urls()
        if urls:
            file_path = urls[0].toLocalFile()
            if file_path.endswith(('.xlsx', '.xls')):
                self.load_excel(file_path)

    def load_excel(self, file_path):
        import pandas as pd  # 懒加载，仅工具箱 Excel→Word 功能触发
        try:
            self.excel_data = pd.read_excel(file_path)
            self.initial_label.hide()
            self.main_widget.show()
            self.update_title_list()
            self.update_preview()
        except Exception as e:
            print(f"Error loading Excel file: {e}")

    def update_title_list(self):
        self.title_list.clear()
        if self.excel_data is not None:
            # 根据当前模式显示可拖拽的项（列名 或 行索引）
            mode = "列模式"
            if hasattr(self, 'mode_combo') and self.mode_combo is not None:
                mode = self.mode_combo.currentText()

            if mode == "列模式":
                columns = self.excel_data.columns.tolist()
                if not columns:
                    QMessageBox.warning(self, "警告", "当前Excel文件没有列标题！")
                    return
                self.title_list.addItems(columns)
                # 标记列表当前展示的是列项
                self.title_list.item_mode = "列模式"
            else:
                rows = [f"行 {i}" for i in range(len(self.excel_data))]
                if not rows:
                    QMessageBox.warning(self, "警告", "当前Excel没有数据行！")
                    return
                self.title_list.addItems(rows)
                # 标记列表当前展示的是行项
                self.title_list.item_mode = "行模式"
        else:
            QMessageBox.warning(self, "警告", "请先加载Excel文件！")

    def update_preview(self):
        self.refresh_preview_area()
        
        if self.excel_data is None:
            return
            
        mode = "列模式"
        if hasattr(self, 'mode_combo') and self.mode_combo is not None:
            mode = self.mode_combo.currentText()
        selected_titles = [self.title_list.item(i).text() for i in range(self.title_list.count())]
        
        # 统一处理：无论是列模式还是行模式，都构建相同的模板结构
        self.preview_columns = OrderedDict()
        for title in selected_titles:
            if mode == "列模式":
                # 列模式：使用列名作为字段标识
                if title in self.excel_data.columns and len(self.excel_data) > 0:
                    value = str(self.excel_data.iloc[0][title])
                else:
                    value = ""
            else:
                # 行模式：使用行索引作为字段标识
                try:
                    if title.startswith("行 "):
                        row_index = int(title.split()[1])
                    else:
                        row_index = int(title)
                    if 0 <= row_index < len(self.excel_data):
                        # 获取该行的所有列数据，用列名作为key
                        row_data = self.excel_data.iloc[row_index]
                        # 取第一列的值作为示例显示
                        first_col = self.excel_data.columns[0] if len(self.excel_data.columns) > 0 else ""
                        value = str(row_data[first_col]) if first_col else ""
                    else:
                        value = ""
                except (ValueError, IndexError):
                    value = ""
            
            item = {"title": title, "value": value, "style": "正文"}
            self.preview_columns[title] = [item]
        
        # 刷新显示（始终以模板 preview_columns 渲染预览区）
        self.refresh_preview_area()

    def on_mode_changed(self, mode):
        # 切换模式时需更新左侧可拖拽项以及预览
        self.update_title_list()
        self.preview_columns.clear()
        self.refresh_preview_area()

    def add_item_to_template(self, title, mode):
        """将拖拽的项目添加到模板"""
        if self.excel_data is None:
            QMessageBox.warning(self, "警告", "请先加载Excel文件！")
            return

        if mode == "列模式":
            # 列模式：使用列名
            if title in self.excel_data.columns:
                try:
                    value = str(self.excel_data.iloc[0][title]) if len(self.excel_data) > 0 else ""
                except Exception:
                    value = "示例值"
            else:
                QMessageBox.warning(self, "警告", f"列 '{title}' 不存在！")
                return
        else:
            # 行模式：使用行索引
            try:
                if title.startswith("行 "):
                    row_index = int(title.split()[1])
                else:
                    row_index = int(title)
                
                if 0 <= row_index < len(self.excel_data):
                    # 获取该行的第一列数据作为示例值
                    first_col = self.excel_data.columns[0] if len(self.excel_data.columns) > 0 else ""
                    value = str(self.excel_data.iloc[row_index][first_col]) if first_col else "示例值"
                else:
                    QMessageBox.warning(self, "警告", f"行索引 {row_index} 超出范围！")
                    return
            except (ValueError, IndexError):
                QMessageBox.warning(self, "警告", f"无法解析行标题: {title}")
                return
        
        # 统一添加到模板
        item = {"title": title, "value": value, "style": "正文"}
        if title not in self.preview_columns:
            self.preview_columns[title] = []
        self.preview_columns[title].append(item)
        self.refresh_preview_area()

    def add_item_to_template_at_position(self, title, mode, position):
        """在指定位置添加项目到模板 - 统一处理"""
        if self.excel_data is None:
            QMessageBox.warning(self, "警告", "请先加载Excel文件！")
            return
            
        # 统一获取数据值逻辑
        if mode == "列模式":
            if title in self.excel_data.columns:
                value = str(self.excel_data.iloc[0][title]) if len(self.excel_data) > 0 else ""
            else:
                QMessageBox.warning(self, "警告", f"列 '{title}' 不存在！")
                return
        else:  # 行模式
            try:
                if title.startswith("行 "):
                    row_index = int(title.split()[1])
                else:
                    row_index = int(title)
                
                if 0 <= row_index < len(self.excel_data):
                    # 获取该行的第一列数据
                    first_col = self.excel_data.columns[0] if len(self.excel_data.columns) > 0 else ""
                    value = str(self.excel_data.iloc[row_index][first_col]) if first_col else ""
                else:
                    QMessageBox.warning(self, "警告", f"行索引 {row_index} 超出范围！")
                    return
            except (ValueError, IndexError):
                QMessageBox.warning(self, "警告", f"无法解析行标题: {title}")
                return
        
        # 统一添加逻辑
        item = {"title": title, "value": value, "style": "正文"}
        # 在preview_columns中找到合适的位置插入
        keys = list(self.preview_columns.keys())
        if position < len(keys):
            target_key = keys[position]
            self.preview_columns[target_key].insert(0, item)  # 插入到该列的开头
        else:
            # 如果位置超出，添加为新列
            self.preview_columns[title] = [item]
        
        self.refresh_preview_area()
        
    def add_item_to_column(self, column_title, title, mode):
        """向指定列添加字段（列内横向排列）"""
        if self.excel_data is None:
            QMessageBox.warning(self, "警告", "请先加载Excel文件！")
            return

        if title in self.excel_data.columns:
            try:
                value = str(self.excel_data.iloc[0][title]) if len(self.excel_data) > 0 else ""
            except Exception:
                value = "示例值"
        else:
            value = "示例值"

        item = {"title": title, "value": value, "style": "正文"}
        if column_title not in self.preview_columns:
            self.preview_columns[column_title] = []
        self.preview_columns[column_title].append(item)
        self.refresh_preview_area()
        
    def add_column_with_item(self, title, mode):
        """在列模式下，新建一列并把字段作为该列的第一个元素"""
        if self.excel_data is None:
            QMessageBox.warning(self, "警告", "请先加载Excel文件！")
            return

        # 获取字段示例值
        if title in self.excel_data.columns:
            try:
                value = str(self.excel_data.iloc[0][title]) if len(self.excel_data) > 0 else ""
            except Exception:
                value = "示例值"
        else:
            value = "示例值"

        item = {"title": title, "value": value, "style": "正文"}
        # 如果列不存在，则创建并加入项目
        if title not in self.preview_columns:
            self.preview_columns[title] = []
        self.preview_columns[title].append(item)
        self.refresh_preview_area()


    def handle_bottom_drop(self, title, mode):
        """处理底部 20px 拖拽区放下事件：列模式新建列，行模式作为普通项加入末尾"""
        if mode == "列模式":
            self.add_column_with_item(title, mode)
        else:
            self.add_item_to_template(title, mode)
    

    def refresh_preview_area(self):
        """刷新预览区域"""
        # 保存当前滚动位置
        scrollbar = self.preview_area.verticalScrollBar()
        saved_scroll_pos = scrollbar.value()
    
        # 清空预览区域 - 修复布局清理逻辑
        while self.preview_layout.count():
            item = self.preview_layout.takeAt(0)
            if item:
                widget = item.widget()
                if widget:
                    widget.setParent(None)
                # 也要清理布局本身
                layout = item.layout()
                if layout:
                    # 递归清理嵌套布局
                    while layout.count():
                        child_item = layout.takeAt(0)
                        if child_item and child_item.widget():
                            child_item.widget().setParent(None)
        
        for col_title, items in self.preview_columns.items():
            col_widget = CompositePreviewItemWidget.ColumnPreviewWidget(col_title, items, self.preview_widget)
            col_widget.add_item_to_column = self.add_item_to_column
            self.preview_layout.addWidget(col_widget)   
            
        self.preview_layout.addStretch()        
        QTimer.singleShot(10, lambda: scrollbar.setValue(saved_scroll_pos))
        
    class DocGenerator(QThread):
        """在后台线程中生成 Word 文档，使用传入的模板快照和遍历方式，避免直接访问 GUI 对象。"""
        finished = Signal(bool, str)  # success, path_or_message

        def __init__(self, preview_columns_snapshot, excel_data, traversal, title, output_path="output.docx"):
            super().__init__()
            self.preview_columns = preview_columns_snapshot
            self.excel_data = excel_data
            self.traversal = traversal  # 'col' or 'row'
            self.title = title
            self.output_path = output_path
            
        
        def _apply_chinese_style(self, doc):
            """
            修改 doc 中常用内置样式的字体为宋体，颜色为黑色。
            """
            from docx.shared import RGBColor, Pt
            from docx.oxml.ns import qn
            from docx.enum.style import WD_STYLE_TYPE
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            styles = doc.styles

            # === 1. 基础段落样式：Normal 已存在，我们只修改它 ===
            normal = styles['Normal']
            normal.font.name = 'SimSun'
            normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            normal.font.color.rgb = RGBColor(0, 0, 0)
            normal.font.size = Pt(10.5)  # 5号字体

            # === 2. 创建自定义标题样式 ===
            # 自定义主标题
            if 'In Title' not in styles:
                title_style = styles.add_style('In Title', WD_STYLE_TYPE.PARAGRAPH)
                title_style.base_style = normal  # 继承 Normal
                title_font = title_style.font
                title_font.name = 'SimSun'
                title_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                title_font.color.rgb = RGBColor(0, 0, 0)
                title_font.size = Pt(16)
                title_font.bold = True
                title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_style.paragraph_format.space_before = Pt(24)
                title_style.paragraph_format.space_after = Pt(18)
                title_style.hidden = False
                title_style.quick_style = True
                title_style.priority = 1
                

            # 自定义一级标题
            if 'In Heading 1' not in styles:
                h1 = styles.add_style('In Heading 1', WD_STYLE_TYPE.PARAGRAPH)
                h1.base_style = normal
                h1.font.name = 'SimSun'
                h1._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                h1.font.color.rgb = RGBColor(0, 0, 0)
                h1.font.size = Pt(14)
                h1.font.bold = True
                h1.paragraph_format.space_before = Pt(18)
                h1.paragraph_format.space_after = Pt(12)
                h1.hidden = False
                h1.quick_style = True
                h1.priority = 1

            # 自定义二级标题
            if 'In Heading 2' not in styles:
                h2 = styles.add_style('In Heading 2', WD_STYLE_TYPE.PARAGRAPH)
                h2.base_style = normal
                h2.font.name = 'SimSun'
                h2._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                h2.font.color.rgb = RGBColor(0, 0, 0)
                h2.font.size = Pt(13)
                h2.font.bold = True
                h2.paragraph_format.space_before = Pt(12)
                h2.paragraph_format.space_after = Pt(6)
                h2.hidden = False
                h2.quick_style = True
                h2.priority = 1

            # 自定义三级标题
            if 'In Heading 3' not in styles:
                h3 = styles.add_style('In Heading 3', WD_STYLE_TYPE.PARAGRAPH)
                h3.base_style = normal
                h3.font.name = 'SimSun'
                h3._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                h3.font.color.rgb = RGBColor(0, 0, 0)
                h3.font.size = Pt(12)
                h3.font.bold = True
                h3.paragraph_format.space_before = Pt(12)
                h3.paragraph_format.space_after = Pt(6)
                h3.hidden = False
                h3.quick_style = True
                h3.priority = 1

            # === 3. 修改其他常用样式（列表、引用等）===
            for style_name in ['Quote', 'List Paragraph', 'List Bullet', 'List Number']:
                if style_name in styles:
                    s = styles[style_name]
                    s.font.name = 'SimSun'
                    s._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                    s.font.color.rgb = RGBColor(0, 0, 0)
                    s.font.size = Pt(10.5)  # 5号字体

            # 字符样式
            for style_name in ['Strong', 'Emphasis']:
                if style_name in styles:
                    s = styles[style_name]
                    s.font.name = 'SimSun'
                    s._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                    s.font.color.rgb = RGBColor(0, 0, 0)
                    s.font.size = Pt(10.5)  # 5号字体
                    
            # === 4. 隐藏 latent styles ===
            latent_styles = doc.styles.latent_styles

            for name in ['Title', 'Heading 1', 'Heading 2', 'Heading 3']:
                if name in latent_styles:
                    ls = latent_styles[name]
                    ls.hidden = True      # 在样式窗格中隐藏
                    ls.quick_style = False  # 不出现在快速样式栏
                    
        
        def run(self):
            from docx import Document  # 懒加载，仅生成文档时触发
            from docx.shared import RGBColor, Pt
            from docx.oxml.ns import qn
            from docx.enum.style import WD_STYLE_TYPE
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            try:
                doc = Document()
                self._apply_chinese_style(doc)

                if self.title:
                    # heading = doc.add_heading(self.title, level=1)
                    heading = doc.add_paragraph(self.title, style='In Title')
                print("开始生成文档...")
                # 如果没有模板则返回空结果
                if not self.preview_columns:
                    doc.save(self.output_path)
                    self.finished.emit(True, self.output_path)
                    return

                # 根据遍历方向生成文档：'col' -> 按行遍历所有行；'row' -> 遍历所有列
                if self.traversal == 'col':
                    row_indices = range(len(self.excel_data)) if self.excel_data is not None else []
                else:
                    row_indices = range(len(self.excel_data)) if self.excel_data is not None else []

                first = True
                for i, row_index in enumerate(row_indices):
                    if not first:
                        doc.add_paragraph("")
                        doc.add_paragraph("-" * 30)
                        doc.add_paragraph("")
                    first = False

                    for col_title, items in self.preview_columns.items():
                        # 构建当前行每个字段的实际值副本
                        items_with_values = []
                        for item in items:
                            field = item.get("title")
                            
                            if self.traversal == 'col':
                                # 列模式：field是列名，获取对应行的数据
                                if self.excel_data is None or field not in self.excel_data.columns:
                                    value = ""
                                else:
                                    try:
                                        value = str(self.excel_data.iloc[row_index][field])
                                        if not value or value.lower() == "nan":
                                            value = ""
                                    except Exception:
                                        value = ""
                            else:
                                # 行模式：field是行索引，获取该行各列的数据
                                try:
                                    if field.startswith("行 "): # 行索引为行号
                                        target_row = int(field.split()[1])
                                    else:
                                        target_row = int(field)
                                    
                                    if 0 <= target_row < len(self.excel_data):
                                        # 获取该行的对应列数据
                                        col_index = row_index  # 使用当前遍历的列索引
                                        if col_index < len(self.excel_data.columns):
                                            col_name = self.excel_data.columns[col_index]
                                            value = str(self.excel_data.iloc[target_row][col_name])
                                            if not value or value.lower() == "nan":
                                                value = ""
                                        else:
                                            value = ""
                                    else:
                                        value = ""
                                except (ValueError, IndexError):
                                    value = ""
                            items_with_values.append({"title": field, "value": value, "style": item.get("style", "正文")})

                        # 使用复合格式将列内所有字段横向合并显示
                        # 采用第一个子项的样式作为整列样式
                        style = items_with_values[0].get("style", "正文") if items_with_values else "正文"
                        display_parts = []
                        for it in items_with_values:
                            v = it.get("value", "")
                            display_parts.append(v if v else "") 
                        display_text = " ".join(display_parts)
                        
                        if display_text == "":
                            continue

                        if style == "一级标题":
                            p = doc.add_paragraph(display_text, style='In Heading 1')
                        elif style == "二级标题":
                            p = doc.add_paragraph(display_text, style='In Heading 2')
                        elif style == "三级标题":
                            p = doc.add_paragraph(display_text, style='In Heading 3')
                        elif style == "加粗":
                            p = doc.add_paragraph()
                            r = p.add_run(display_text)
                            r.style = "Strong"
                        elif style == "斜体":
                            p = doc.add_paragraph()
                            r = p.add_run(display_text)
                            r.style = "Emphasis"
                        elif style == "引用":
                            p = doc.add_paragraph(display_text, style='Quote')
                        elif style == "无序列表":
                            p = doc.add_paragraph(style='List Bullet')
                            p.add_run(display_text)
                        elif style == "有序列表":
                            p = doc.add_paragraph(style='List Number')
                            p.add_run(display_text)
                        else:
                            doc.add_paragraph(display_text, style='Normal')

                doc.save(self.output_path)
                self.finished.emit(True, self.output_path)
            except Exception as e:
                self.finished.emit(False, str(e))

    def generate_word(self):
        """生成Word文档 - 批量处理所有数据行"""
        if self.excel_data is None:
            QMessageBox.warning(self, "警告", "请先加载Excel文件！")
            return

        # 统一使用 preview_columns 作为模板模型；如果没有模板则提示
        if not self.preview_columns:
            QMessageBox.warning(self, "警告", "请先完成模板！")
            return

        # 禁用生成按钮并显示不可取消的进度对话框（指示正在后台生成）
        try:
            self.generate_btn.setEnabled(False)
            self.progress_dialog = QProgressDialog("正在生成Word...", "None", 0, 0, self)
            self.progress_dialog.setWindowModality(Qt.WindowModal)
            self.progress_dialog.setCancelButton(None)
            self.progress_dialog.setMinimumDuration(0)
            self.progress_dialog.show()

            if not self.preview_columns:
                QMessageBox.warning(self, "警告", "请先完成模板！")
                self.progress_dialog.close()
                self.generate_btn.setEnabled(True)
                return

            output_path = "output.docx"
            # 在主线程对模板和选择行进行快照，避免线程访问 GUI 对象
            preview_columns_snapshot = deepcopy(self.preview_columns)
            traversal = 'col' if (hasattr(self, 'mode_combo') and self.mode_combo.currentText() == '列模式') else 'row'
            title_snapshot = self.primary_title.text().strip() if hasattr(self, 'primary_title') else ''

            self._doc_worker = ExcelToWordView.DocGenerator(preview_columns_snapshot, self.excel_data, traversal, title_snapshot, output_path)
            self._doc_worker.finished.connect(self._on_doc_generation_finished)
            self._doc_worker.start()
        except Exception as e:
            QMessageBox.critical(self, "错误", f"启动后台生成失败：{str(e)}")

    


    def _on_doc_generation_finished(self, success, path_or_msg):
        # 在主线程恢复 UI 状态并提示结果
        try:
            if hasattr(self, 'progress_dialog') and self.progress_dialog is not None:
                self.progress_dialog.close()
        except Exception:
            pass
        try:
            self.generate_btn.setEnabled(True)
        except Exception:
            pass
        if success:
            QMessageBox.information(self, "成功", f"Word文档已生成：{path_or_msg}")
        else:
            QMessageBox.critical(self, "错误", f"生成Word文档失败：{path_or_msg}")

    def preview_drag_enter(self, event):
        """预览区域拖拽进入事件"""
        if event.mimeData().hasFormat("application/x-etoword-item"):
            event.acceptProposedAction()

    def preview_drop_event(self, event):
        """预览区域拖拽放下事件"""
        if event.mimeData().hasFormat("application/x-etoword-item"):
            data = event.mimeData().data("application/x-etoword-item").data().decode()
            title, mode = data.split("|")
            
            # 如果是来自左侧的项目，添加到末尾
            if mode in ["列模式", "行模式"]:
                self.add_item_to_template(title, mode)
            
            event.acceptProposedAction()
    
if __name__ == "__main__":
    app = QApplication(sys.argv)
    view = ExcelToWordView()
    view.show()
    sys.exit(app.exec())
import pandas as pd
from docx import Document
from typing import List, Dict, Optional, Union
from pathlib import Path

class ExcelToWordHandler:
    """Excel转Word处理器"""
    
    def __init__(self):
        self.data: Optional[pd.DataFrame] = None
        self.file_path: Optional[str] = None
    
    def load_excel(self, file_path: str) -> bool:
        """
        加载Excel文件
        
        Args:
            file_path: Excel文件路径
            
        Returns:
            bool: 加载是否成功
        """
        try:
            self.file_path = file_path
            self.data = pd.read_excel(file_path)
            return True
        except Exception as e:
            print(f"加载Excel文件失败: {e}")
            return False
    
    def get_columns(self) -> List[str]:
        """获取列标题列表"""
        if self.data is not None:
            return self.data.columns.tolist()
        return []
    
    def get_rows_count(self) -> int:
        """获取行数"""
        if self.data is not None:
            return len(self.data)
        return 0
    
    def get_column_preview_data(self, column_name: str, row_index: int = 0) -> str:
        """
        获取指定列的预览数据
        
        Args:
            column_name: 列名
            row_index: 行索引，默认为0（第一行数据）
            
        Returns:
            str: 数据值的字符串表示
        """
        if self.data is not None and column_name in self.data.columns:
            try:
                value = self.data.iloc[row_index][column_name]
                return str(value) if pd.notna(value) else ""
            except (IndexError, KeyError):
                return ""
        return ""
    
    def get_row_preview_data(self, row_index: int) -> List[str]:
        """
        获取指定行的所有列数据
        
        Args:
            row_index: 行索引
            
        Returns:
            List[str]: 该行所有列的数据列表
        """
        if self.data is not None and 0 <= row_index < len(self.data):
            try:
                row_data = self.data.iloc[row_index]
                return [str(val) if pd.notna(val) else "" for val in row_data]
            except IndexError:
                return []
        return []
    
    def generate_word_document(self, 
                             primary_title: str = "",
                             preview_items: Optional[List[Dict[str, str]]] = None,
                             footer_content: str = "",
                             output_path: str = "output.docx") -> bool:
        """
        生成Word文档
        
        Args:
            primary_title: 一级标题
            preview_items: 预览项列表 [{title: str, value: str, style: str}]
            footer_content: 末尾内容
            output_path: 输出文件路径
            
        Returns:
            bool: 生成是否成功
        """
        try:
            doc = Document()
            
            # 添加一级标题
            if primary_title.strip():
                doc.add_heading(primary_title.strip(), level=1)
            
            # 添加预览内容（应用样式）
            if preview_items:
                for item in preview_items:
                    title = item.get("title", "")
                    value = item.get("value", "")
                    style = item.get("style", "默认")
                    
                    if not value.strip():
                        continue
                        
                    if style == "二级标题":
                        doc.add_heading(value.strip(), level=2)
                    elif style == "无序列表":
                        p = doc.add_paragraph()
                        p.style = 'List Bullet'
                        p.add_run(value.strip())
                    elif style == "有序列表":
                        p = doc.add_paragraph()
                        p.style = 'List Number'
                        p.add_run(value.strip())
                    else:  # 默认样式
                        # 如果有标题，显示为"标题: 值"格式
                        display_text = f"{title}: {value}" if title else value
                        doc.add_paragraph(display_text.strip())
            
            # 添加末尾内容（用分割线分隔）
            if footer_content.strip():
                doc.add_paragraph("----------", style="Intense Quote")
                doc.add_paragraph(footer_content.strip())
            
            # 保存文档
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"生成Word文档失败: {e}")
            return False

    def generate_word_from_preview(self,
                                   excel_df: Optional[pd.DataFrame],
                                   primary_title: str,
                                   preview_columns: Optional[Dict[str, List[Dict[str,str]]]] = None,
                                   preview_items: Optional[List[Dict[str,str]]] = None,
                                   mode: str = "列模式",
                                   output_path: str = "output.docx",
                                   progress_callback=None) -> bool:
        """
        根据视图中的预览数据生成 Word 文档，支持列模式和行模式，并通过回调上报进度。

        Args:
            excel_df: 原始 DataFrame（用于按行/列取值）
            primary_title: 一级标题
            preview_columns: 列模式下的列定义（OrderedDict 或 dict: 列名 -> list of item dicts）
            preview_items: 行模式下的预览项列表
            mode: "列模式" 或 "行模式"
            output_path: 输出文件路径
            progress_callback: 可选回调，签名 progress_callback(percent: int)

        Returns:
            bool: 成功与否
        """
        try:
            doc = Document()

            # 添加一级标题
            if primary_title and primary_title.strip():
                doc.add_heading(primary_title.strip(), level=1)

            # 处理不同模式
            if mode == "列模式":
                if excel_df is None or len(excel_df) == 0 or not preview_columns:
                    # nothing to do
                    doc.save(output_path)
                    if progress_callback:
                        progress_callback(100)
                    return True

                total_rows = len(excel_df)
                for row_index in range(total_rows):
                    # 行分隔
                    if row_index > 0:
                        doc.add_paragraph("")
                        doc.add_paragraph("-" * 30)
                        doc.add_paragraph("")

                    # 跳过当行中任一模板字段无数据的整行（用户要求不生成“无数据”文本）
                    skip_row = False
                    for col_title, items in (preview_columns.items() if hasattr(preview_columns, 'items') else preview_columns.items()):
                        for item in items:
                            field = (item.get("title") or "")
                            if not field or field not in excel_df.columns:
                                skip_row = True
                                break
                            # 使用 Series.get 避免类型提示问题
                            cell_val = excel_df.iloc[row_index].get(field, None)
                            if cell_val is None or (isinstance(cell_val, float) and pd.isna(cell_val)) or (isinstance(cell_val, str) and cell_val.strip() == ""):
                                skip_row = True
                                break
                        if skip_row:
                            break

                    if skip_row:
                        # 按行回调进度并跳过此行
                        if progress_callback:
                            pct = int((row_index + 1) * 100 / total_rows)
                            progress_callback(min(pct, 100))
                        continue

                    # 遍历每列并生成内容
                    for col_title, items in (preview_columns.items() if hasattr(preview_columns, 'items') else preview_columns.items()):
                        items_with_values = []
                        for item in items:
                            field = (item.get("title") or "")
                            # 此处已保证字段存在且不为空
                            cell_val = excel_df.iloc[row_index].get(field, None)
                            value = "(空)" if (cell_val is None or (isinstance(cell_val, float) and pd.isna(cell_val))) else str(cell_val)
                            items_with_values.append({"title": field, "value": value, "style": item.get("style", "默认")})

                        style = items_with_values[0].get("style", "默认") if items_with_values else "默认"
                        # 组合显示
                        display_parts = [it.get("value", "(空)") or "(空)" for it in items_with_values]
                        display_text = " ".join(display_parts)
                        if style == "二级标题":
                            doc.add_heading(display_text, level=2)
                        elif style == "无序列表":
                            p = doc.add_paragraph(style='List Bullet')
                            p.add_run(display_text)
                        elif style == "有序列表":
                            p = doc.add_paragraph(style='List Number')
                            p.add_run(display_text)
                        else:
                            doc.add_paragraph(display_text)

                    # 进度回调（按行）
                    if progress_callback:
                        pct = int((row_index + 1) * 100 / total_rows)
                        progress_callback(min(pct, 100))

            else:  # 行模式
                if excel_df is None or preview_items is None:
                    doc.save(output_path)
                    if progress_callback:
                        progress_callback(100)
                    return True

                # 先解析出要处理的行索引列表
                selected_rows = []
                for item in preview_items:
                    try:
                        if item.get("title", "").startswith("Row "):
                            row_index = int(item["title"].split()[1])
                        else:
                            row_index = int(item.get("title", ""))
                        if 0 <= row_index < len(excel_df):
                            selected_rows.append((row_index, item))
                    except Exception:
                        continue

                total = len(selected_rows)
                for i, (row_index, item) in enumerate(selected_rows): #/ 遍历每行并生成内容
                    if i > 0:
                        doc.add_paragraph("")
                        doc.add_paragraph("-" * 30)
                        doc.add_paragraph("")

                    row_data = excel_df.iloc[row_index]
                    row_values = [str(val) if not pd.isna(val) else "" for val in row_data]
                    if "".join(row_values) == "": # 跳过无数据行
                        continue
                    for j, (column_name, value) in enumerate(zip(excel_df.columns, row_values)):
                        style = item.get("style", "默认")
                        # 值展示不带标题前缀
                        if style == "二级标题":
                            doc.add_heading(value, level=2)
                        elif style == "无序列表":
                            p = doc.add_paragraph(style='List Bullet')
                            p.add_run(value)
                        elif style == "有序列表":
                            p = doc.add_paragraph(style='List Number')
                            p.add_run(value)
                        else:
                            doc.add_paragraph(value)

                    if progress_callback and total > 0:
                        pct = int((i + 1) * 100 / total)
                        progress_callback(min(pct, 100))

            # 保存并回调完成
            doc.save(output_path)
            if progress_callback:
                progress_callback(100)
            return True

        except Exception as e:
            print(f"生成Word文档失败: {e}")
            return False
    
    def get_file_info(self) -> Dict[str, Union[str, int]]:
        """获取文件基本信息"""
        info = {
            "file_path": self.file_path or "",
            "rows": self.get_rows_count(),
            "columns": len(self.get_columns()),
            "column_names": self.get_columns()
        }
        return info
